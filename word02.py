from docx import Document  # 导入python-docx库，用于读取Word文档
import time               # 导入time库，用于计时

def wordFile(docx_path):
    """
    解析 Word 文档，返回带表格的层次结构字典。
    每个节点包含：content（正文列表）、subsections（子标题）、tables（表格列表）。
    """
    # 判断段落样式名是否为“Heading X”，如果是则返回标题级别数字，否则返回None。
    def get_heading_level(style_name):
        if style_name and style_name.startswith('Heading'):
            try:
                return int(style_name.split(' ')[1])
            except Exception:
                return None
        return None

    doc = Document(docx_path)
    # root为根节点，存储全文结构。
    root = {"content": [], "subsections": {}, "tables": []}
    # node_stack为当前目录层级的栈。
    node_stack = [root]
    # 遍历段落和表格。
    para_idx, table_idx = 0, 0
    # 导入底层类型，便于判断元素类型。
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    for child in doc.element.body:
        if isinstance(child, CT_P):
            para = doc.paragraphs[para_idx]
            para_idx += 1
            text = para.text.strip()
            if not text:
                continue
            level = get_heading_level(para.style.name)
            if level:
                while len(node_stack) > level:
                    node_stack.pop()
                parent = node_stack[-1]["subsections"]
                parent[text] = {"content": [], "subsections": {}, "tables": []}
                node_stack.append(parent[text])
            else:
                node_stack[-1]["content"].append(text)
        elif isinstance(child, CT_Tbl):
            table = doc.tables[table_idx]
            table_idx += 1
            table_content = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            node_stack[-1]["tables"].append(table_content)
    def clean(node):
        node["content"] = [c for c in node["content"] if c]
        for sub in node["subsections"].values():
            clean(sub)
    clean(root)
    return root

def tables(node, indent=0):
    """
    递归打印正文、表格和标题，保持层次结构。
    """
    prefix = "  " * indent  # 生成缩进前缀（每层多两个空格）
    for content in node["content"]:
        print(f"{prefix}-  {content}")  # 打印正文内容，带缩进
    for table in node.get("tables", []):
        for row in table:
            # 打印表格一行，空单元格用'-'占位
            row_text = [str(cell) if cell else '-' for cell in row]
            print(" | ".join(row_text))
    for title, subnode in node["subsections"].items():
        print(f"{prefix}- {title}")  # 打印标题，带缩进
        tables(subnode, indent + 1)  # 递归打印子节点，缩进加深

if __name__ == "__main__":
    docx_path = "01.docx"  # Word文档路径
    start = time.time()    # 记录开始时间
    outline = wordFile(docx_path)  # 解析文档，获得结构树
    print("文档：")
    tables(outline)        # 打印正文、表格和标题
    print(f"\n处理时间: {time.time() - start:.2f} 秒")  # 打印处理用时