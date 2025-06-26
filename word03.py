from docx import Document

def tables(doc_path):
    doc = Document(doc_path)
    if not doc.tables:
        return None
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() if cell.text else '-' for cell in row.cells]
            print(" | ".join(row_text))

if __name__ == "__main__":
    tables("01.docx")