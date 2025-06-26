from docx import Document

doc = Document()

doc.add_heading('复杂表格示例', level=1)
doc.add_paragraph('下面是一个包含多行多列、跨行跨列合并的复杂表格：')

# 创建8行6列的表格
table = doc.add_table(rows=8, cols=6)
table.style = 'Table Grid'

# 第一行：大表头，全部合并
table.cell(0, 0).merge(table.cell(0, 5))
table.cell(0, 0).text = '学生成绩总表'

# 第二行：分组表头，前两列合并，后四列分两组
table.cell(1, 0).merge(table.cell(1, 1))
table.cell(1, 0).text = '基本信息'
table.cell(1, 2).merge(table.cell(1, 3))
table.cell(1, 2).text = '语文成绩'
table.cell(1, 4).merge(table.cell(1, 5))
table.cell(1, 4).text = '数学成绩'

# 第三行：详细表头
table.cell(2, 0).text = '姓名'
table.cell(2, 1).text = '性别'
table.cell(2, 2).text = '期中'
table.cell(2, 3).text = '期末'
table.cell(2, 4).text = '期中'
table.cell(2, 5).text = '期末'

# 第四~八行：学生数据，部分单元格合并
students = [
    ['张三', '男', '85', '88', '90', '92'],
    ['李四', '女', '78', '80', '85', '87'],
    ['王五', '男', '90', '93', '95', '97'],
    ['赵六', '女', '82', '85', '88', '90'],
    ['钱七', '男', '88', '90', '91', '94']
]

for i, stu in enumerate(students):
    row = 3 + i
    for j, val in enumerate(stu):
        table.cell(row, j).text = val

# 第八行：备注，前五列合并
table.cell(8-1, 0).merge(table.cell(8-1, 4))
table.cell(8-1, 0).text = '备注：以上数据仅供参考'
table.cell(8-1, 5).text = ''

doc.save('complex_table.docx')
print('complex_table.docx 已生成')